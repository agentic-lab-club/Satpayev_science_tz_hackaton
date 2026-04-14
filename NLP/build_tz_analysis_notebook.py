from __future__ import annotations

import json
import textwrap
from pathlib import Path


def md_cell(source: str) -> dict:
    return {
        "cell_type": "markdown",
        "metadata": {},
        "source": textwrap.dedent(source).strip("\n").splitlines(keepends=True),
    }


def code_cell(source: str) -> dict:
    return {
        "cell_type": "code",
        "execution_count": None,
        "metadata": {},
        "outputs": [],
        "source": textwrap.dedent(source).strip("\n").splitlines(keepends=True),
    }


NOTEBOOK_CELLS = [
    md_cell(
        """
        # TZ Analysis MVP

        Интерактивный Jupyter Notebook для анализа и улучшения технических заданий.

        Pipeline:
        `file -> parsing -> structure -> semantic -> score -> issues -> improved TZ -> extraction -> confirmation -> final report`
        """
    ),
    code_cell(
        r'''
        # === CELL: 1. Установка зависимостей ===
        %pip install -q python-docx pdfplumber pytesseract openai pydantic
        '''
    ),
    md_cell(
        """
        ## 2. Импорты и базовые настройки

        Здесь задаются:
        - импорты;
        - логирование;
        - API-настройки для LLM;
        - обязательные разделы ТЗ;
        - demo-файл, чтобы notebook запускался сверху вниз без дополнительной подготовки.
        """
    ),
    code_cell(
        r'''
        # === CELL: 2. Импорты и базовые настройки ===
        import json
        import logging
        import os
        import re
        import textwrap
        import unicodedata
        from collections import defaultdict
        from functools import lru_cache
        from pathlib import Path
        from typing import Dict, List, Optional

        from IPython.display import Markdown, display
        from pydantic import BaseModel, Field, ValidationError

        try:
            import pdfplumber
        except ImportError:
            pdfplumber = None

        try:
            import pytesseract
        except ImportError:
            pytesseract = None

        try:
            from docx import Document as DocxDocument
        except ImportError:
            DocxDocument = None

        try:
            from openai import OpenAI
        except ImportError:
            OpenAI = None

        logging.basicConfig(
            level=logging.INFO,
            format="%(asctime)s | %(levelname)s | %(message)s",
        )
        logger = logging.getLogger("tz_notebook")

        ROOT_DIR = Path.cwd()
        OUTPUT_DIR = ROOT_DIR / "notebook_outputs"
        OUTPUT_DIR.mkdir(exist_ok=True)

        OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "").strip()
        OPENAI_BASE_URL = os.getenv("OPENAI_BASE_URL", "").strip() or None
        OPENAI_MODEL = os.getenv("OPENAI_MODEL", "").strip() or "gpt-4o-mini"

        GROQ_API_KEY = os.getenv("GROQ_API_KEY", "").strip()
        if not OPENAI_API_KEY and GROQ_API_KEY:
            OPENAI_API_KEY = GROQ_API_KEY
            OPENAI_BASE_URL = OPENAI_BASE_URL or "https://api.groq.com/openai/v1"
            OPENAI_MODEL = os.getenv("GROQ_MODEL", "").strip() or "llama-3.1-70b-versatile"

        MANDATORY_SECTIONS: Dict[str, List[str]] = {
            "цель": ["цель", "цели", "purpose", "objective"],
            "задачи": ["задачи", "задача", "scope", "объем работ"],
            "требования": ["требования", "функциональные требования", "нефункциональные требования"],
            "сроки": ["сроки", "дедлайн", "этапы", "timeline", "календарный план"],
            "kpi": ["kpi", "метрики", "показатели эффективности", "ключевые показатели"],
            "ожидаемые результаты": ["ожидаемые результаты", "результаты", "deliverables", "итог"],
        }

        VAGUE_PATTERNS = {
            "как можно скорее": "Укажите конкретный дедлайн или срок в днях/неделях.",
            "в кратчайшие сроки": "Замените на точный срок выполнения.",
            "при необходимости": "Опишите критерии, когда именно возникает необходимость.",
            "удобный": "Опишите измеримые критерии удобства или UX-метрики.",
            "современный": "Уточните конкретные требования к стилю, технологиям или интерфейсу.",
            "быстро": "Укажите допустимое время выполнения или SLA.",
            "качественный": "Добавьте измеримые критерии качества.",
            "эффективный": "Опишите, как именно измеряется эффективность.",
            "и т.д.": "Перечислите полный список элементов без сокращения 'и т.д.'.",
        }

        CONTRADICTION_PAIRS = [
            ("онлайн", "офлайн"),
            ("автоматически", "вручную"),
            ("не требуется", "обязательно"),
            ("без api", "через api"),
            ("не позднее", "после"),
        ]

        DEMO_FILE_PATH = ROOT_DIR / "demo_tz_sample.txt"
        if not DEMO_FILE_PATH.exists():
            DEMO_FILE_PATH.write_text(
                textwrap.dedent(
                    """
                    1. Цель проекта
                    Разработать MVP-систему анализа технических заданий для научных и инновационных проектов в формате Jupyter Notebook.

                    2. Задачи
                    - Реализовать загрузку документов в форматах PDF, DOCX и TXT.
                    - Выполнить анализ структуры и смысла документа.
                    - Подготовить улучшенную версию ТЗ.

                    3. Требования
                    - Интерфейс должен быть удобным и современным.
                    - Система должна быстро обрабатывать документы.
                    - Необходимо поддержать импорт PDF, DOCX и TXT.
                    - При необходимости допускается ручная корректировка результатов.

                    4. Сроки
                    - Первый прототип подготовить до 15.06.2026.
                    - Финальную версию выпустить в течение 2 недель после согласования замечаний.

                    5. Ожидаемые результаты
                    - Рабочий notebook с полным pipeline анализа.
                    - Список найденных проблем и рекомендаций.
                    - Улучшенный текст ТЗ.
                    """
                ).strip(),
                encoding="utf-8",
            )

        print("ROOT_DIR:", ROOT_DIR)
        print("OUTPUT_DIR:", OUTPUT_DIR)
        print("DEMO_FILE_PATH:", DEMO_FILE_PATH)
        print("LLM enabled:", bool(OPENAI_API_KEY and OpenAI))
        print("Model:", OPENAI_MODEL if OPENAI_API_KEY else "heuristic mode")
        '''
    ),
    md_cell(
        """
        ## 3. Pydantic-схемы

        Все ответы от LLM и внутренние артефакты прогоняются через Pydantic.
        """
    ),
    code_cell(
        r'''
        # === CELL: 3. Pydantic-схемы ===
        class Section(BaseModel):
            title: str
            content: str
            length: int


        class StructureResult(BaseModel):
            sections: List[Section] = Field(default_factory=list)
            found_sections: List[str] = Field(default_factory=list)
            missing_sections: List[str] = Field(default_factory=list)
            weak_sections: List[str] = Field(default_factory=list)
            empty_sections: List[str] = Field(default_factory=list)
            matched_required_sections: Dict[str, str] = Field(default_factory=dict)


        class IssueItem(BaseModel):
            issue_type: str = ""
            section: str = ""
            quote: str = ""
            explanation: str = ""
            recommendation: str = ""


        class IssueList(BaseModel):
            items: List[IssueItem] = Field(default_factory=list)


        class ExtractionItem(BaseModel):
            value: str
            section: str = ""
            confidence: float = Field(default=0.0, ge=0.0, le=1.0)


        class ExtractionList(BaseModel):
            items: List[ExtractionItem] = Field(default_factory=list)


        class SemanticAnalysisResult(BaseModel):
            ambiguities: List[IssueItem] = Field(default_factory=list)
            contradictions: List[IssueItem] = Field(default_factory=list)
            missing_elements: List[IssueItem] = Field(default_factory=list)
            requirements: List[ExtractionItem] = Field(default_factory=list)
            deadlines: List[ExtractionItem] = Field(default_factory=list)
            kpis: List[ExtractionItem] = Field(default_factory=list)
            expected_results: List[ExtractionItem] = Field(default_factory=list)


        class ScoreResult(BaseModel):
            total_score: int
            breakdown: Dict[str, int]
            explanation: List[str] = Field(default_factory=list)


        class ImprovedTZResult(BaseModel):
            summary_of_changes: List[str] = Field(default_factory=list)
            improved_text: str
        '''
    ),
    md_cell(
        """
        ## 4. Загрузка и парсинг документа

        Поддерживаются:
        - `TXT` через стандартное чтение;
        - `DOCX` через `python-docx`;
        - `PDF` через `pdfplumber`;
        - OCR для PDF включается только при доступности `pytesseract`.
        """
    ),
    code_cell(
        r'''
        # === CELL: 4. Загрузка и парсинг документа ===
        def normalize_text(text: str) -> str:
            if text is None:
                return ""
            text = unicodedata.normalize("NFKC", text)
            text = text.replace("\xa0", " ").replace("\u200b", " ")
            text = text.replace("\r\n", "\n").replace("\r", "\n").replace("\t", " ")
            text = re.sub(r"[ ]{2,}", " ", text)
            text = re.sub(r"\n[ \t]+", "\n", text)
            text = re.sub(r"\n{3,}", "\n\n", text)
            return text.strip()


        def ensure_non_empty_text(text: str) -> str:
            cleaned = normalize_text(text)
            if not cleaned:
                raise ValueError("Документ пустой или не удалось извлечь текст.")
            return cleaned


        def _ocr_pdf_page(page) -> str:
            if pytesseract is None:
                return ""
            try:
                page_image = page.to_image(resolution=200)
                pil_image = page_image.original
                ocr_text = pytesseract.image_to_string(pil_image, lang="rus+eng")
                return normalize_text(ocr_text)
            except Exception as exc:
                logger.warning("OCR fallback failed: %s", exc)
                return ""


        def load_txt(file_path: Path) -> str:
            return ensure_non_empty_text(file_path.read_text(encoding="utf-8", errors="ignore"))


        def load_docx(file_path: Path) -> str:
            if DocxDocument is None:
                raise ImportError("python-docx не установлен.")
            document = DocxDocument(file_path)
            paragraphs = [p.text.strip() for p in document.paragraphs if p.text and p.text.strip()]
            table_chunks = []
            for table in document.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        table_chunks.append(row_text)
            raw_text = "\n".join(paragraphs + table_chunks)
            return ensure_non_empty_text(raw_text)


        def load_pdf(file_path: Path, use_ocr_if_needed: bool = True) -> str:
            if pdfplumber is None:
                raise ImportError("pdfplumber не установлен.")

            pages = []
            with pdfplumber.open(file_path) as pdf:
                for idx, page in enumerate(pdf.pages, start=1):
                    page_text = normalize_text(page.extract_text() or "")
                    if not page_text and use_ocr_if_needed:
                        page_text = _ocr_pdf_page(page)
                    if page_text:
                        pages.append(f"[PAGE {idx}]\n{page_text}")

            return ensure_non_empty_text("\n\n".join(pages))


        def load_document(file_path: str) -> str:
            path = Path(file_path)
            if not path.exists():
                raise FileNotFoundError(f"Файл не найден: {path.resolve()}")

            suffix = path.suffix.lower()
            logger.info("Loading document: %s", path.name)

            if suffix == ".txt":
                return load_txt(path)
            if suffix == ".docx":
                return load_docx(path)
            if suffix == ".pdf":
                return load_pdf(path)

            raise ValueError(f"Неподдерживаемый формат файла: {suffix}")


        def preview_text(text: str, max_chars: int = 700) -> str:
            text = normalize_text(text)
            if len(text) <= max_chars:
                return text
            return text[:max_chars].rstrip() + "..."
        '''
    ),
    md_cell(
        """
        ## 5. Выделение структуры документа

        Ноутбук ищет заголовки по нумерации, CAPS и ключевым словам, затем сверяет документ с обязательной структурой ТЗ.
        """
    ),
    code_cell(
        r'''
        # === CELL: 5. Выделение структуры документа ===
        HEADING_PATTERNS = [
            re.compile(r"^\d+(\.\d+)*[.)]?\s+[A-Za-zА-Яа-яЁё]"),
            re.compile(r"^(раздел|section)\s+\d+", re.IGNORECASE),
        ]


        def canonicalize_label(value: str) -> str:
            return normalize_text(value).lower()


        def is_heading(line: str) -> bool:
            line = normalize_text(line)
            if not line or len(line) > 140:
                return False

            for pattern in HEADING_PATTERNS:
                if pattern.search(line):
                    return True

            normalized = canonicalize_label(line)
            if any(normalized.startswith(alias) for aliases in MANDATORY_SECTIONS.values() for alias in aliases):
                return True

            letters = [char for char in line if char.isalpha()]
            if letters:
                uppercase_ratio = sum(char.isupper() for char in letters) / len(letters)
                if uppercase_ratio >= 0.75 and len(line.split()) <= 10:
                    return True

            return False


        def split_into_sections(text: str) -> List[Section]:
            lines = [line.strip() for line in normalize_text(text).split("\n")]
            sections: List[Section] = []
            current_title = "Общий контекст"
            current_content: List[str] = []

            for line in lines:
                if not line:
                    if current_content and current_content[-1] != "":
                        current_content.append("")
                    continue

                if is_heading(line):
                    content = normalize_text("\n".join(current_content))
                    if current_title != "Общий контекст" or content:
                        sections.append(
                            Section(title=current_title, content=content, length=len(content))
                        )
                    current_title = line
                    current_content = []
                else:
                    current_content.append(line)

            content = normalize_text("\n".join(current_content))
            if current_title or content:
                sections.append(Section(title=current_title, content=content, length=len(content)))

            sections = [section for section in sections if section.title or section.content]
            if not sections:
                sections = [Section(title="Общий контекст", content=normalize_text(text), length=len(normalize_text(text)))]
            return sections


        def match_required_sections(sections: List[Section]) -> Dict[str, str]:
            matched = {}
            for canonical, aliases in MANDATORY_SECTIONS.items():
                for section in sections:
                    haystack = f"{section.title}\n{section.content[:500]}".lower()
                    if any(alias in haystack for alias in aliases):
                        matched[canonical] = section.title
                        break
            return matched


        def evaluate_section_quality(sections: List[Section], matched_required: Dict[str, str]) -> Dict[str, List[str]]:
            weak_sections = []
            empty_sections = []

            for canonical, actual_title in matched_required.items():
                section = next((item for item in sections if item.title == actual_title), None)
                if section is None:
                    continue

                word_count = len(section.content.split())
                if not section.content.strip():
                    empty_sections.append(canonical)
                elif word_count < 20 or section.length < 120:
                    weak_sections.append(canonical)

            return {
                "weak_sections": weak_sections,
                "empty_sections": empty_sections,
            }


        def extract_sections(text: str) -> StructureResult:
            cleaned = ensure_non_empty_text(text)
            sections = split_into_sections(cleaned)
            matched_required = match_required_sections(sections)
            found_sections = list(matched_required.keys())
            missing_sections = [item for item in MANDATORY_SECTIONS if item not in matched_required]
            quality = evaluate_section_quality(sections, matched_required)

            return StructureResult(
                sections=sections,
                found_sections=found_sections,
                missing_sections=missing_sections,
                weak_sections=quality["weak_sections"],
                empty_sections=quality["empty_sections"],
                matched_required_sections=matched_required,
            )


        def print_structure_summary(result: StructureResult) -> None:
            print("Найденные разделы:")
            for section in result.sections:
                print(f"- {section.title} | length={section.length}")

            print("\nОбязательные разделы:")
            print("  found   :", ", ".join(result.found_sections) or "—")
            print("  missing :", ", ".join(result.missing_sections) or "—")
            print("  weak    :", ", ".join(result.weak_sections) or "—")
            print("  empty   :", ", ".join(result.empty_sections) or "—")
        '''
    ),
    md_cell(
        """
        ## 6. LLM-слой и безопасный JSON-парсинг

        Каждый смысловой шаг выполняется отдельным запросом.
        Если API-ключ недоступен, notebook автоматически переключается в heuristic mode.
        """
    ),
    code_cell(
        r'''
        # === CELL: 6. LLM-слой и безопасный JSON-парсинг ===
        def llm_enabled() -> bool:
            return bool(OPENAI_API_KEY and OpenAI is not None)


        @lru_cache(maxsize=1)
        def get_llm_client():
            if not llm_enabled():
                return None
            return OpenAI(api_key=OPENAI_API_KEY, base_url=OPENAI_BASE_URL)


        def strip_code_fences(text: str) -> str:
            text = text.strip()
            text = re.sub(r"^```json\s*", "", text)
            text = re.sub(r"^```\s*", "", text)
            text = re.sub(r"\s*```$", "", text)
            return text.strip()


        def extract_json_blob(text: str) -> str:
            text = strip_code_fences(text)
            start = text.find("{")
            end = text.rfind("}")
            if start == -1 or end == -1 or start >= end:
                raise ValueError("JSON object not found in LLM response.")
            return text[start : end + 1]


        def truncate_for_llm(text: str, max_chars: int = 9000) -> str:
            text = normalize_text(text)
            if len(text) <= max_chars:
                return text
            return text[:max_chars].rstrip() + "\n\n[TRUNCATED]"


        def llm_call(task_instruction: str, input_text: str, response_model):
            if not llm_enabled():
                raise RuntimeError("LLM is disabled. Heuristic fallback will be used instead.")

            schema_hint = {
                IssueList: '{"items":[{"issue_type":"","section":"","quote":"","explanation":"","recommendation":""}]}',
                ExtractionList: '{"items":[{"value":"","section":"","confidence":0.0}]}',
                ImprovedTZResult: '{"summary_of_changes":[""],"improved_text":""}',
            }.get(response_model, '{"items":[]}')

            messages = [
                {
                    "role": "system",
                    "content": (
                        "Ты анализируешь технические задания. "
                        "Отвечай только валидным JSON без markdown и пояснений."
                    ),
                },
                {
                    "role": "user",
                    "content": (
                        f"{task_instruction}\n\n"
                        f"Верни JSON строго такого формата:\n{schema_hint}\n\n"
                        f"Документ:\n{truncate_for_llm(input_text)}"
                    ),
                },
            ]

            client = get_llm_client()
            try:
                response = client.chat.completions.create(
                    model=OPENAI_MODEL,
                    messages=messages,
                    temperature=0,
                    response_format={"type": "json_object"},
                )
            except Exception:
                response = client.chat.completions.create(
                    model=OPENAI_MODEL,
                    messages=messages,
                    temperature=0,
                )

            content = response.choices[0].message.content or "{}"
            payload = json.loads(extract_json_blob(content))
            return response_model.model_validate(payload)
        '''
    ),
    md_cell(
        """
        ## 7. Смысловой анализ

        Функции:
        - `detect_ambiguities`
        - `detect_contradictions`
        - `detect_missing_elements`
        - `extract_requirements`
        - `extract_deadlines`
        - `extract_kpis`
        - `extract_expected_results`
        """
    ),
    code_cell(
        r'''
        # === CELL: 7. Смысловой анализ ===
        DATE_PATTERN = re.compile(
            r"\b\d{1,2}[./]\d{1,2}[./]\d{2,4}\b|"
            r"\b\d{1,2}\s+(?:января|февраля|марта|апреля|мая|июня|июля|августа|сентября|октября|ноября|декабря)\s+\d{4}\b|"
            r"\b\d+\s+(?:рабочих\s+)?(?:дней|недель|месяцев)\b",
            re.IGNORECASE,
        )


        def sentence_candidates(text: str) -> List[str]:
            text = normalize_text(text)
            chunks = re.split(r"(?<=[.!?])\s+|\n+", text)
            cleaned = []
            for chunk in chunks:
                candidate = chunk.strip(" -•")
                if not candidate:
                    continue
                if is_heading(candidate):
                    continue
                cleaned.append(candidate)
            return cleaned


        def locate_section(snippet: str, structure_result: StructureResult) -> str:
            snippet_norm = canonicalize_label(snippet)[:120]
            for section in structure_result.sections:
                haystack = canonicalize_label(section.content)
                if snippet_norm and snippet_norm in haystack:
                    return section.title
            return ""


        def deduplicate_issues(items: List[IssueItem]) -> List[IssueItem]:
            result = []
            seen = set()
            for item in items:
                key = (
                    item.issue_type.strip().lower(),
                    item.section.strip().lower(),
                    item.quote.strip().lower(),
                    item.explanation.strip().lower(),
                )
                if key in seen:
                    continue
                seen.add(key)
                result.append(item)
            return result


        def deduplicate_extractions(items: List[ExtractionItem]) -> List[ExtractionItem]:
            result = []
            seen = set()
            for item in sorted(items, key=lambda x: x.confidence, reverse=True):
                key = (item.value.strip().lower(), item.section.strip().lower())
                if key in seen:
                    continue
                seen.add(key)
                result.append(item)
            return result


        def get_focus_text(structure_result: StructureResult, aliases: List[str], fallback_text: str, max_chars: int = 5000) -> str:
            chunks = []
            alias_set = [alias.lower() for alias in aliases]
            for section in structure_result.sections:
                haystack = f"{section.title}\n{section.content}".lower()
                if any(alias in haystack for alias in alias_set):
                    chunks.append(f"{section.title}\n{section.content}")

            joined = "\n\n".join(chunks) if chunks else fallback_text
            return truncate_for_llm(joined, max_chars=max_chars)


        def heuristic_detect_ambiguities(text: str, structure_result: StructureResult) -> IssueList:
            issues = []
            sentences = sentence_candidates(text)
            for sentence in sentences:
                lower_sentence = sentence.lower()
                for phrase, recommendation in VAGUE_PATTERNS.items():
                    if phrase in lower_sentence:
                        issues.append(
                            IssueItem(
                                issue_type="ambiguity",
                                section=locate_section(sentence, structure_result),
                                quote=sentence[:240],
                                explanation=f"Найдена размытая формулировка: '{phrase}'.",
                                recommendation=recommendation,
                            )
                        )
            return IssueList(items=deduplicate_issues(issues))


        def heuristic_detect_contradictions(text: str, structure_result: StructureResult) -> IssueList:
            issues = []
            for sentence in sentence_candidates(text):
                lower_sentence = sentence.lower()
                for left, right in CONTRADICTION_PAIRS:
                    if left in lower_sentence and right in lower_sentence:
                        issues.append(
                            IssueItem(
                                issue_type="contradiction",
                                section=locate_section(sentence, structure_result),
                                quote=sentence[:240],
                                explanation=f"В одной формулировке встречаются потенциально конфликтующие условия: '{left}' и '{right}'.",
                                recommendation="Разделите условия или оставьте только один согласованный вариант.",
                            )
                        )
            return IssueList(items=deduplicate_issues(issues))


        def heuristic_detect_missing_elements(text: str, structure_result: StructureResult) -> IssueList:
            issues = []
            for section_name in structure_result.missing_sections:
                issues.append(
                    IssueItem(
                        issue_type="missing_element",
                        section=section_name,
                        quote="",
                        explanation=f"Обязательный раздел '{section_name}' не найден в документе.",
                        recommendation=f"Добавьте раздел '{section_name}' с конкретными и измеримыми формулировками.",
                    )
                )

            for section_name in structure_result.weak_sections:
                issues.append(
                    IssueItem(
                        issue_type="missing_element",
                        section=section_name,
                        quote="",
                        explanation=f"Раздел '{section_name}' найден, но выглядит слишком коротким или поверхностным.",
                        recommendation=f"Раскройте раздел '{section_name}' подробнее: критерии, ограничения, числовые значения, результаты.",
                    )
                )
            return IssueList(items=deduplicate_issues(issues))


        def heuristic_extract_requirements(text: str, structure_result: StructureResult) -> ExtractionList:
            patterns = ["должен", "должна", "должны", "требуется", "необходимо", "обязан", "поддерж", "обеспеч"]
            items = []
            focus_text = get_focus_text(structure_result, MANDATORY_SECTIONS["требования"], text)
            for sentence in sentence_candidates(focus_text):
                lower_sentence = sentence.lower()
                if any(pattern in lower_sentence for pattern in patterns):
                    confidence = 0.85 if any(token in lower_sentence for token in ["долж", "обязан"]) else 0.7
                    items.append(
                        ExtractionItem(
                            value=sentence[:300],
                            section=locate_section(sentence, structure_result),
                            confidence=confidence,
                        )
                    )
            return ExtractionList(items=deduplicate_extractions(items))


        def heuristic_extract_deadlines(text: str, structure_result: StructureResult) -> ExtractionList:
            items = []
            focus_text = get_focus_text(structure_result, MANDATORY_SECTIONS["сроки"], text)
            for sentence in sentence_candidates(focus_text):
                matches = DATE_PATTERN.findall(sentence)
                if matches or any(token in sentence.lower() for token in ["срок", "дедлайн", "этап", "до ", "не позднее"]):
                    value = sentence[:300]
                    confidence = 0.9 if matches else 0.65
                    items.append(
                        ExtractionItem(
                            value=value,
                            section=locate_section(sentence, structure_result),
                            confidence=confidence,
                        )
                    )
            return ExtractionList(items=deduplicate_extractions(items))


        def heuristic_extract_kpis(text: str, structure_result: StructureResult) -> ExtractionList:
            items = []
            focus_text = get_focus_text(structure_result, MANDATORY_SECTIONS["kpi"], text)
            patterns = ["kpi", "%", "не менее", "не более", "точность", "скорость", "время отклика", "снизить", "увеличить"]
            for sentence in sentence_candidates(focus_text):
                lower_sentence = sentence.lower()
                if any(pattern in lower_sentence for pattern in patterns) and re.search(r"\d|%", sentence):
                    items.append(
                        ExtractionItem(
                            value=sentence[:300],
                            section=locate_section(sentence, structure_result),
                            confidence=0.8,
                        )
                    )
            return ExtractionList(items=deduplicate_extractions(items))


        def heuristic_extract_expected_results(text: str, structure_result: StructureResult) -> ExtractionList:
            items = []
            patterns = ["ожида", "результат", "итог", "deliverable", "будет подготовлен", "прототип", "отчет", "отчёт", "ноутбук", "notebook"]

            matched_titles = [
                title for canonical, title in structure_result.matched_required_sections.items()
                if canonical == "ожидаемые результаты"
            ]
            for title in matched_titles:
                section = next((item for item in structure_result.sections if item.title == title), None)
                if not section:
                    continue
                for line in section.content.splitlines():
                    candidate = line.strip(" -•")
                    if len(candidate) < 10:
                        continue
                    items.append(
                        ExtractionItem(
                            value=candidate[:300],
                            section=section.title,
                            confidence=0.82,
                        )
                    )

            focus_text = get_focus_text(structure_result, MANDATORY_SECTIONS["ожидаемые результаты"], text)
            for sentence in sentence_candidates(focus_text):
                lower_sentence = sentence.lower()
                if any(pattern in lower_sentence for pattern in patterns):
                    items.append(
                        ExtractionItem(
                            value=sentence[:300],
                            section=locate_section(sentence, structure_result),
                            confidence=0.72,
                        )
                    )
            return ExtractionList(items=deduplicate_extractions(items))


        def call_issue_detector(task_name: str, instruction: str, focus_text: str, fallback_fn, structure_result: StructureResult) -> IssueList:
            if not llm_enabled():
                return fallback_fn(focus_text, structure_result)
            try:
                result = llm_call(instruction, focus_text, IssueList)
                return IssueList(items=deduplicate_issues(result.items))
            except (ValidationError, ValueError, RuntimeError, Exception) as exc:
                logger.warning("%s failed, fallback to heuristics: %s", task_name, exc)
                return fallback_fn(focus_text, structure_result)


        def call_extractor(task_name: str, instruction: str, focus_text: str, fallback_fn, structure_result: StructureResult) -> ExtractionList:
            if not llm_enabled():
                return fallback_fn(focus_text, structure_result)
            try:
                result = llm_call(instruction, focus_text, ExtractionList)
                return ExtractionList(items=deduplicate_extractions(result.items))
            except (ValidationError, ValueError, RuntimeError, Exception) as exc:
                logger.warning("%s failed, fallback to heuristics: %s", task_name, exc)
                return fallback_fn(focus_text, structure_result)


        def detect_ambiguities(text: str, structure_result: StructureResult) -> dict:
            focus_text = truncate_for_llm(text)
            instruction = (
                "Найди до 7 неоднозначных, размытых или неизмеримых формулировок в ТЗ. "
                "Для каждой укажи issue_type='ambiguity', section, точную quote, explanation и recommendation."
            )
            result = call_issue_detector(
                "detect_ambiguities",
                instruction,
                focus_text,
                heuristic_detect_ambiguities,
                structure_result,
            )
            return result.model_dump()


        def detect_contradictions(text: str, structure_result: StructureResult) -> dict:
            focus_text = truncate_for_llm(text)
            instruction = (
                "Найди до 5 противоречий или потенциально конфликтующих требований в ТЗ. "
                "Для каждой проблемы верни issue_type='contradiction', section, quote, explanation и recommendation."
            )
            result = call_issue_detector(
                "detect_contradictions",
                instruction,
                focus_text,
                heuristic_detect_contradictions,
                structure_result,
            )
            return result.model_dump()


        def detect_missing_elements(text: str, structure_result: StructureResult) -> dict:
            focus_text = (
                f"Обязательные разделы: {', '.join(MANDATORY_SECTIONS.keys())}\n"
                f"Найденные разделы: {', '.join(structure_result.found_sections)}\n"
                f"Отсутствуют: {', '.join(structure_result.missing_sections)}\n\n"
                f"{truncate_for_llm(text)}"
            )
            instruction = (
                "Найди отсутствующие или недостаточно раскрытые элементы ТЗ. "
                "Для каждой проблемы верни issue_type='missing_element', section, quote, explanation и recommendation."
            )
            result = call_issue_detector(
                "detect_missing_elements",
                instruction,
                focus_text,
                heuristic_detect_missing_elements,
                structure_result,
            )
            return result.model_dump()


        def extract_requirements(text: str, structure_result: StructureResult) -> dict:
            focus_text = get_focus_text(structure_result, MANDATORY_SECTIONS["требования"], text)
            instruction = (
                "Извлеки требования из ТЗ. "
                "Верни JSON со списком items, где у каждого элемента есть value, section и confidence от 0.0 до 1.0."
            )
            result = call_extractor(
                "extract_requirements",
                instruction,
                focus_text,
                heuristic_extract_requirements,
                structure_result,
            )
            return result.model_dump()


        def extract_deadlines(text: str, structure_result: StructureResult) -> dict:
            focus_text = get_focus_text(structure_result, MANDATORY_SECTIONS["сроки"], text)
            instruction = (
                "Извлеки сроки, даты, этапы и дедлайны из ТЗ. "
                "Верни JSON со списком items: value, section, confidence."
            )
            result = call_extractor(
                "extract_deadlines",
                instruction,
                focus_text,
                heuristic_extract_deadlines,
                structure_result,
            )
            return result.model_dump()


        def extract_kpis(text: str, structure_result: StructureResult) -> dict:
            focus_text = get_focus_text(structure_result, MANDATORY_SECTIONS["kpi"], text)
            instruction = (
                "Извлеки KPI, метрики и измеримые показатели успеха из ТЗ. "
                "Верни JSON со списком items: value, section, confidence."
            )
            result = call_extractor(
                "extract_kpis",
                instruction,
                focus_text,
                heuristic_extract_kpis,
                structure_result,
            )
            return result.model_dump()


        def extract_expected_results(text: str, structure_result: StructureResult) -> dict:
            focus_text = get_focus_text(structure_result, MANDATORY_SECTIONS["ожидаемые результаты"], text)
            instruction = (
                "Извлеки ожидаемые результаты, deliverables и конечные артефакты из ТЗ. "
                "Верни JSON со списком items: value, section, confidence."
            )
            result = call_extractor(
                "extract_expected_results",
                instruction,
                focus_text,
                heuristic_extract_expected_results,
                structure_result,
            )
            return result.model_dump()


        def run_semantic_analysis(text: str, structure_result: StructureResult) -> SemanticAnalysisResult:
            ambiguities = IssueList.model_validate(detect_ambiguities(text, structure_result))
            contradictions = IssueList.model_validate(detect_contradictions(text, structure_result))
            missing_elements = IssueList.model_validate(detect_missing_elements(text, structure_result))
            requirements = ExtractionList.model_validate(extract_requirements(text, structure_result))
            deadlines = ExtractionList.model_validate(extract_deadlines(text, structure_result))
            kpis = ExtractionList.model_validate(extract_kpis(text, structure_result))
            expected_results = ExtractionList.model_validate(extract_expected_results(text, structure_result))

            return SemanticAnalysisResult(
                ambiguities=ambiguities.items,
                contradictions=contradictions.items,
                missing_elements=missing_elements.items,
                requirements=requirements.items,
                deadlines=deadlines.items,
                kpis=kpis.items,
                expected_results=expected_results.items,
            )
        '''
    ),
    md_cell(
        """
        ## 8. Explainable scoring и генерация улучшений

        Score разбит на 6 прозрачных компонентов.
        """
    ),
    code_cell(
        r'''
        # === CELL: 8. Explainable scoring и генерация улучшений ===
        def clamp(value: int, low: int, high: int) -> int:
            return max(low, min(high, int(value)))


        def calculate_score(structure_result: StructureResult, semantic_result: SemanticAnalysisResult) -> ScoreResult:
            required_count = len(MANDATORY_SECTIONS)
            found_count = len(structure_result.found_sections)

            structure_base = round(25 * found_count / required_count)
            structure_penalty = 2 * len(structure_result.weak_sections) + 4 * len(structure_result.empty_sections)
            structure_score = clamp(structure_base - structure_penalty, 0, 25)

            completeness_hits = 0
            for key in ["цель", "задачи", "требования", "сроки"]:
                if key in structure_result.found_sections and key not in structure_result.weak_sections:
                    completeness_hits += 1
            if semantic_result.deadlines:
                completeness_hits += 0.5
            if semantic_result.expected_results:
                completeness_hits += 0.5
            completeness_score = clamp(round(20 * completeness_hits / 5), 0, 20)

            requirements_score = 0
            if "требования" in structure_result.found_sections:
                requirements_score += 8
            requirements_score += min(12, len(semantic_result.requirements) * 3)
            requirements_score = clamp(requirements_score, 0, 20)

            consistency_penalty = min(15, len(semantic_result.contradictions) * 4 + min(6, len(semantic_result.ambiguities)))
            consistency_score = clamp(15 - consistency_penalty, 0, 15)

            kpi_score = 0
            if "kpi" in structure_result.found_sections:
                kpi_score += 4
            kpi_score += min(6, len(semantic_result.kpis) * 3)
            kpi_score = clamp(kpi_score, 0, 10)

            results_score = 0
            if "ожидаемые результаты" in structure_result.found_sections:
                results_score += 4
            results_score += min(6, len(semantic_result.expected_results) * 2)
            results_score = clamp(results_score, 0, 10)

            breakdown = {
                "structure_25": structure_score,
                "completeness_20": completeness_score,
                "requirements_20": requirements_score,
                "consistency_15": consistency_score,
                "kpi_10": kpi_score,
                "expected_results_10": results_score,
            }
            total_score = sum(breakdown.values())

            explanation = [
                f"Структура: найдено {found_count}/{required_count} обязательных разделов.",
                f"Полнота: слабые разделы -> {', '.join(structure_result.weak_sections) or 'нет'}.",
                f"Требования: извлечено {len(semantic_result.requirements)} формулировок.",
                f"Непротиворечивость: ambiguities={len(semantic_result.ambiguities)}, contradictions={len(semantic_result.contradictions)}.",
                f"KPI: найдено {len(semantic_result.kpis)} элементов.",
                f"Ожидаемые результаты: найдено {len(semantic_result.expected_results)} элементов.",
            ]

            return ScoreResult(
                total_score=clamp(total_score, 0, 100),
                breakdown=breakdown,
                explanation=explanation,
            )


        def collect_all_issues(semantic_result: SemanticAnalysisResult) -> List[IssueItem]:
            return deduplicate_issues(
                semantic_result.ambiguities
                + semantic_result.contradictions
                + semantic_result.missing_elements
            )


        def generate_recommendations(issues: List[IssueItem]) -> List[str]:
            recommendations = []
            seen = set()
            for issue in issues:
                text = issue.recommendation.strip()
                if not text:
                    continue
                key = text.lower()
                if key in seen:
                    continue
                seen.add(key)
                recommendations.append(text)
            return recommendations


        def build_missing_section_template(section_name: str) -> str:
            templates = {
                "цель": "Цель проекта: [УТОЧНИТЬ конечный бизнес/исследовательский эффект и критерий успеха].",
                "задачи": "- Подзадача 1: [УТОЧНИТЬ]\n- Подзадача 2: [УТОЧНИТЬ]\n- Подзадача 3: [УТОЧНИТЬ]",
                "требования": "- Функциональные требования: [УТОЧНИТЬ]\n- Нефункциональные требования: [УТОЧНИТЬ]\n- Ограничения: [УТОЧНИТЬ]",
                "сроки": "- Этап 1: [ДАТА]\n- Этап 2: [ДАТА]\n- Финальный дедлайн: [ДАТА]",
                "kpi": "- KPI 1: [МЕТРИКА И ЦЕЛЕВОЕ ЗНАЧЕНИЕ]\n- KPI 2: [МЕТРИКА И ЦЕЛЕВОЕ ЗНАЧЕНИЕ]",
                "ожидаемые результаты": "- Артефакт 1: [УТОЧНИТЬ]\n- Артефакт 2: [УТОЧНИТЬ]\n- Критерий приемки: [УТОЧНИТЬ]",
            }
            return templates.get(section_name, "[УТОЧНИТЬ]")


        def heuristic_generate_improved_tz(original_text: str, issues: List[IssueItem], structure_result: StructureResult) -> ImprovedTZResult:
            summary_of_changes = []
            recommendations = generate_recommendations(issues)

            for section_name in structure_result.missing_sections:
                summary_of_changes.append(f"Добавлен каркас раздела '{section_name}'.")

            for issue in issues:
                if issue.issue_type == "ambiguity" and issue.quote:
                    summary_of_changes.append(f"Помечена размытая формулировка для уточнения: '{issue.quote[:80]}'.")

            appended_sections = []
            for section_name in structure_result.missing_sections:
                appended_sections.append(
                    f"{section_name.title()}\n{build_missing_section_template(section_name)}"
                )

            clarification_lines = []
            for issue in issues:
                if issue.quote and issue.recommendation:
                    clarification_lines.append(
                        f"- Было: {issue.quote}\n  Стало: {issue.recommendation}"
                    )

            improved_parts = [normalize_text(original_text)]

            if appended_sections:
                improved_parts.append(
                    "Дополнительно включить в ТЗ следующие разделы:\n\n" + "\n\n".join(appended_sections)
                )

            if clarification_lines:
                improved_parts.append(
                    "Уточнения по формулировкам:\n" + "\n".join(clarification_lines[:10])
                )

            if recommendations:
                improved_parts.append(
                    "Ключевые рекомендации:\n" + "\n".join(f"- {item}" for item in recommendations[:10])
                )

            improved_text = "\n\n".join(part for part in improved_parts if part.strip())
            return ImprovedTZResult(
                summary_of_changes=deduplicate_plain_list(summary_of_changes) or ["Добавлены подсказки по усилению ТЗ."],
                improved_text=improved_text,
            )


        def deduplicate_plain_list(items: List[str]) -> List[str]:
            result = []
            seen = set()
            for item in items:
                key = item.strip().lower()
                if not key or key in seen:
                    continue
                seen.add(key)
                result.append(item.strip())
            return result


        def generate_improved_tz(original_text: str, issues: List[IssueItem], structure_result: StructureResult) -> ImprovedTZResult:
            if not llm_enabled():
                return heuristic_generate_improved_tz(original_text, issues, structure_result)

            issue_summary = [
                {
                    "issue_type": issue.issue_type,
                    "section": issue.section,
                    "quote": issue.quote,
                    "recommendation": issue.recommendation,
                }
                for issue in issues[:15]
            ]

            instruction = (
                "Подготовь улучшенную версию ТЗ. "
                "Сохрани исходный смысл, но усили формулировки, добавь недостающие элементы, "
                "а если точных данных нет, пометь место как [УТОЧНИТЬ]. "
                f"Проблемы для исправления: {json.dumps(issue_summary, ensure_ascii=False)}"
            )

            try:
                return llm_call(instruction, original_text, ImprovedTZResult)
            except (ValidationError, ValueError, RuntimeError, Exception) as exc:
                logger.warning("generate_improved_tz failed, fallback to heuristics: %s", exc)
                return heuristic_generate_improved_tz(original_text, issues, structure_result)
        '''
    ),
    md_cell(
        """
        ## 9. Извлечение сущностей и подтверждение пользователем

        Пользователь может:
        - `y` — подтвердить;
        - `e` — отредактировать;
        - `n` — отклонить.
        """
    ),
    code_cell(
        r'''
        # === CELL: 9. Извлечение сущностей и подтверждение пользователем ===
        def build_entity_bundle(semantic_result: SemanticAnalysisResult) -> Dict[str, List[ExtractionItem]]:
            return {
                "requirements": semantic_result.requirements,
                "deadlines": semantic_result.deadlines,
                "kpis": semantic_result.kpis,
                "expected_results": semantic_result.expected_results,
            }


        def print_entity_bundle(entity_bundle: Dict[str, List[ExtractionItem]]) -> None:
            for category, items in entity_bundle.items():
                print(f"\n{category.upper()}:")
                if not items:
                    print("- ничего не найдено")
                    continue
                for idx, item in enumerate(items, start=1):
                    print(f"{idx}. {item.value} | section={item.section or '—'} | confidence={item.confidence:.2f}")


        def confirm_entities_interactively(
            entity_bundle: Dict[str, List[ExtractionItem]],
            interactive: bool = True,
            auto_confirm_threshold: float = 0.60,
        ) -> Dict[str, List[dict]]:
            confirmed: Dict[str, List[dict]] = defaultdict(list)

            for category, items in entity_bundle.items():
                if not items:
                    continue

                for item in items:
                    if not interactive:
                        if item.confidence >= auto_confirm_threshold:
                            confirmed[category].append(item.model_dump())
                        continue

                    print(f"\n[{category}] {item.value}")
                    print(f"section={item.section or '—'} | confidence={item.confidence:.2f}")
                    while True:
                        action = input("[y] подтвердить / [e] редактировать / [n] отклонить: ").strip().lower()
                        if action in {"y", "e", "n"}:
                            break
                        print("Введите y, e или n.")

                    if action == "y":
                        confirmed[category].append(item.model_dump())
                    elif action == "e":
                        edited_value = input("Новое значение: ").strip()
                        if edited_value:
                            confirmed[category].append(
                                ExtractionItem(
                                    value=edited_value,
                                    section=item.section,
                                    confidence=max(item.confidence, 0.99),
                                ).model_dump()
                            )

            return dict(confirmed)
        '''
    ),
    md_cell(
        """
        ## 10. Финальный pipeline и отчёт

        Ноутбук собирает единый отчёт:
        - структура;
        - проблемы;
        - explainable score;
        - рекомендации;
        - улучшенная версия;
        - подтверждённые сущности.
        """
    ),
    code_cell(
        r'''
        # === CELL: 10. Финальный pipeline и отчёт ===
        def format_issues_for_display(issues: List[IssueItem]) -> str:
            if not issues:
                return "Проблемы не найдены."
            lines = []
            for idx, issue in enumerate(issues, start=1):
                lines.append(
                    f"{idx}. [{issue.issue_type}] section={issue.section or '—'}\n"
                    f"   quote: {issue.quote or '—'}\n"
                    f"   explanation: {issue.explanation or '—'}\n"
                    f"   recommendation: {issue.recommendation or '—'}"
                )
            return "\n".join(lines)


        def save_report_as_json(report: dict, file_name: str = "tz_analysis_report.json") -> Path:
            output_path = OUTPUT_DIR / file_name
            with output_path.open("w", encoding="utf-8") as f:
                json.dump(report, f, ensure_ascii=False, indent=2)
            return output_path


        def render_final_report(report: dict) -> None:
            structure_result = StructureResult.model_validate(report["structure"])
            semantic_result = SemanticAnalysisResult.model_validate(report["semantic"])
            score_result = ScoreResult.model_validate(report["score"])
            improved_result = ImprovedTZResult.model_validate(report["improved_tz"])

            display(Markdown("## Финальный отчёт"))

            print("1) Структура документа")
            print_structure_summary(structure_result)

            print("\n2) Найденные проблемы")
            print(format_issues_for_display(collect_all_issues(semantic_result)))

            print("\n3) Explainable score")
            print("total_score:", score_result.total_score)
            print("breakdown:", json.dumps(score_result.breakdown, ensure_ascii=False, indent=2))
            print("explanation:")
            for item in score_result.explanation:
                print("-", item)

            print("\n4) Рекомендации")
            recommendations = report["recommendations"]
            if recommendations:
                for item in recommendations:
                    print("-", item)
            else:
                print("- рекомендаций нет")

            print("\n5) Улучшенная версия ТЗ")
            print("summary_of_changes:")
            for item in improved_result.summary_of_changes:
                print("-", item)
            print("\nimproved_text:\n")
            print(improved_result.improved_text)

            print("\n6) Подтверждённые сущности")
            print(json.dumps(report["confirmed_entities"], ensure_ascii=False, indent=2))


        def run_pipeline(file_path: str, interactive_confirmation: bool = False) -> dict:
            raw_text = load_document(file_path)
            structure_result = extract_sections(raw_text)
            semantic_result = run_semantic_analysis(raw_text, structure_result)
            score_result = calculate_score(structure_result, semantic_result)
            all_issues = collect_all_issues(semantic_result)
            recommendations = generate_recommendations(all_issues)
            improved_result = generate_improved_tz(raw_text, all_issues, structure_result)
            entity_bundle = build_entity_bundle(semantic_result)
            confirmed_entities = confirm_entities_interactively(
                entity_bundle,
                interactive=interactive_confirmation,
            )

            report = {
                "file_path": str(file_path),
                "raw_text_preview": preview_text(raw_text),
                "structure": structure_result.model_dump(),
                "semantic": semantic_result.model_dump(),
                "score": score_result.model_dump(),
                "recommendations": recommendations,
                "improved_tz": improved_result.model_dump(),
                "extracted_entities": {
                    key: [item.model_dump() for item in values]
                    for key, values in entity_bundle.items()
                },
                "confirmed_entities": confirmed_entities,
            }

            report_path = save_report_as_json(report)
            print(f"JSON report saved to: {report_path}")
            render_final_report(report)
            return report
        '''
    ),
    md_cell(
        """
        ## 11. Запуск

        По умолчанию используется demo-файл.
        Замените путь на любой `PDF`, `DOCX` или `TXT`.
        """
    ),
    code_cell(
        r'''
        # === CELL: 11. Запуск ===
        FILE_PATH = str(DEMO_FILE_PATH)
        INTERACTIVE_CONFIRMATION = False  # Поставьте True в Jupyter, чтобы подтверждать сущности через input()

        analysis_report = run_pipeline(
            file_path=FILE_PATH,
            interactive_confirmation=INTERACTIVE_CONFIRMATION,
        )
        '''
    ),
]


def build_notebook() -> dict:
    return {
        "cells": NOTEBOOK_CELLS,
        "metadata": {
            "kernelspec": {
                "display_name": "Python 3",
                "language": "python",
                "name": "python3",
            },
            "language_info": {
                "name": "python",
                "version": "3.11",
            },
        },
        "nbformat": 4,
        "nbformat_minor": 5,
    }


def main() -> None:
    output_path = Path("tz_analysis_mvp.ipynb")
    notebook = build_notebook()
    output_path.write_text(json.dumps(notebook, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"Notebook created: {output_path.resolve()}")


if __name__ == "__main__":
    main()
