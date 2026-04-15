from __future__ import annotations

import base64
import io
import re
from dataclasses import dataclass
from pathlib import Path
from urllib.parse import urlparse

import httpx
from docx import Document
from pypdf import PdfReader

from app.core.config import get_settings
from app.core.errors import AIServiceError
from app.schemas.analyze import AnalyzeRequest
from app.schemas.common import DocumentMetadata


@dataclass(slots=True)
class ParsedDocument:
    metadata: DocumentMetadata
    text: str
    source: str
    truncated: bool = False


class ParserService:
    """Document parser for text, DOCX, and PDF inputs."""

    def parse(self, request: AnalyzeRequest) -> ParsedDocument:
        settings = get_settings()
        text, source, parser = self._load_document(request)

        truncated = False
        if len(text) > settings.max_input_chars:
            text = text[: settings.max_input_chars]
            truncated = True

        metadata = DocumentMetadata(
            file_path=request.file_url or request.filename,
            filename=request.filename,
            content_type=request.content_type,
            character_count=len(text),
            parser=f"{parser}{'_truncated' if truncated else ''}",
        )
        return ParsedDocument(metadata=metadata, text=text, source=source, truncated=truncated)

    def get_document_metadata(self, request: AnalyzeRequest) -> DocumentMetadata:
        return self.parse(request).metadata

    def _load_document(self, request: AnalyzeRequest) -> tuple[str, str, str]:
        if request.text and request.text.strip():
            return request.text.strip(), "text", "plain_text"
        if request.file_base64 and request.file_base64.strip():
            raw_bytes = self._decode_base64(request.file_base64)
            return self._extract_from_bytes(
                raw_bytes=raw_bytes,
                filename=request.filename,
                content_type=request.content_type,
            )
        if request.file_url and request.file_url.strip():
            raw_bytes, source = self._read_from_file_url(request.file_url.strip())
            text, source_kind, parser = self._extract_from_bytes(
                raw_bytes=raw_bytes,
                filename=request.filename,
                content_type=request.content_type,
            )
            return text, source or source_kind, parser
        raise AIServiceError(code="invalid_input", message="Either text, file_base64, or file_url is required")

    def _decode_base64(self, payload: str) -> bytes:
        try:
            return base64.b64decode(payload, validate=False)
        except Exception as exc:  # pragma: no cover - handled as user input error
            raise AIServiceError(code="invalid_file", message="file_base64 is not valid base64") from exc

    def _read_from_file_url(self, file_url: str) -> tuple[bytes, str]:
        parsed = urlparse(file_url)
        if parsed.scheme in {"http", "https"}:
            with httpx.Client(timeout=get_settings().request_timeout_seconds) as client:
                response = client.get(file_url)
                response.raise_for_status()
                return response.content, "http_url"

        if parsed.scheme == "file":
            local_path = Path(parsed.path)
            return self._read_local_file(local_path), "file_url"

        local_path = Path(file_url)
        if local_path.is_absolute() and local_path.exists():
            return self._read_local_file(local_path), "local_path"

        candidate_roots = [
            Path.cwd(),
            Path(__file__).resolve().parents[2],
            Path(__file__).resolve().parents[3],
        ]
        candidate_names = [local_path]
        if not local_path.parts or len(local_path.parts) == 1:
            candidate_names.extend(
                [
                    Path("docs") / local_path.name,
                    Path("docs/spikes") / local_path.name,
                ]
            )

        for root in candidate_roots:
            for candidate_name in candidate_names:
                candidate = root / candidate_name
                if candidate.exists():
                    return self._read_local_file(candidate), "local_path"

        raise AIServiceError(code="file_not_found", message=f"Unable to read file_url: {file_url}")

    def _read_local_file(self, path: Path) -> bytes:
        try:
            return path.read_bytes()
        except Exception as exc:  # pragma: no cover - handled as user input error
            raise AIServiceError(code="file_read_failed", message=f"Unable to read local file: {path}") from exc

    def _extract_from_bytes(self, raw_bytes: bytes, filename: str, content_type: str) -> tuple[str, str, str]:
        filename_lower = filename.lower()
        content_type_lower = content_type.lower()

        if filename_lower.endswith(".docx") or "word" in content_type_lower:
            return self._extract_from_docx(raw_bytes), "docx", "docx"
        if filename_lower.endswith(".pdf") or "pdf" in content_type_lower:
            return self._extract_from_pdf(raw_bytes), "pdf", "pdf"

        try:
            return raw_bytes.decode("utf-8"), "txt", "utf8_text"
        except UnicodeDecodeError:
            return raw_bytes.decode("utf-8", errors="ignore"), "txt", "utf8_text_lossy"

    def _extract_from_docx(self, raw_bytes: bytes) -> str:
        document = Document(io.BytesIO(raw_bytes))
        parts: list[str] = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts).strip()

    def _extract_from_pdf(self, raw_bytes: bytes) -> str:
        reader = PdfReader(io.BytesIO(raw_bytes))
        pages: list[str] = []
        for page in reader.pages:
            text = page.extract_text() or ""
            if text.strip():
                pages.append(text.strip())
        combined = "\n".join(pages).strip()
        if combined:
            return combined
        return self._extract_pdf_via_fallback(raw_bytes)

    def _extract_pdf_via_fallback(self, raw_bytes: bytes) -> str:
        try:
            import pdfplumber
        except Exception:  # pragma: no cover - dependency issue fallback
            return ""

        extracted: list[str] = []
        with pdfplumber.open(io.BytesIO(raw_bytes)) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ""
                if text.strip():
                    extracted.append(text.strip())
        return "\n".join(extracted).strip()

    def _split_lines(self, text: str) -> list[str]:
        return [line.strip() for line in re.split(r"\r?\n", text) if line.strip()]


def get_parser_service() -> ParserService:
    return ParserService()
