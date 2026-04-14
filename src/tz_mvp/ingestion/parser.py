"""File ingestion and text extraction for notebook workflow."""

from __future__ import annotations

import io
import logging
from pathlib import Path
from typing import Optional

from docx import Document as DocxDocument
from pypdf import PdfReader

from tz_mvp.core.config import TZAnalysisConfig
from tz_mvp.core.models import Document
from tz_mvp.utils.text import normalize_text

LOGGER = logging.getLogger(__name__)


class DocumentIngestor:
    """Loads source files and converts them into normalized text documents."""

    def __init__(self, config: TZAnalysisConfig) -> None:
        self.config = config

    def load_path(self, path: str | Path) -> Document:
        """Loads a document from local file path."""

        file_path = Path(path)
        suffix = file_path.suffix.lower()
        raw_bytes = file_path.read_bytes()
        return self.load_bytes(file_path.name, raw_bytes, file_type=suffix.lstrip("."))

    def load_bytes(self, filename: str, data: bytes, file_type: Optional[str] = None) -> Document:
        """Loads a document from in-memory bytes, used by FileUpload widget."""

        suffix = (file_type or Path(filename).suffix.lstrip(".")).lower()
        document = Document(original_filename=filename, file_type=suffix or "txt")
        try:
            if suffix == "txt":
                raw_text = data.decode("utf-8", errors="ignore")
            elif suffix == "docx":
                raw_text = self._extract_docx(data)
            elif suffix == "pdf":
                raw_text = self._extract_pdf(data)
                if not raw_text.strip() and self.config.ocr_enabled:
                    raw_text = self._ocr_pdf(data)
            else:
                raise ValueError(f"Unsupported file type: {suffix}")
            document.raw_text = raw_text
            document.normalized_text = normalize_text(raw_text)
            document.parse_status = "success" if document.normalized_text else "empty"
        except Exception as exc:  # pragma: no cover - defensive notebook path
            LOGGER.exception("Failed to ingest %s", filename)
            document.parse_status = f"error: {exc}"
            document.raw_text = ""
            document.normalized_text = ""
        return document

    def summarize(self, document: Document) -> dict:
        """Builds short summary for notebook preview."""

        text = document.normalized_text
        preview = text[:500] + ("..." if len(text) > 500 else "")
        return {
            "document_id": document.id,
            "filename": document.original_filename,
            "file_type": document.file_type,
            "parse_status": document.parse_status,
            "char_count": len(text),
            "word_count": len(text.split()),
            "line_count": len(text.splitlines()),
            "preview": preview,
        }

    def _extract_docx(self, data: bytes) -> str:
        stream = io.BytesIO(data)
        doc = DocxDocument(stream)
        paragraphs = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
        table_chunks = []
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    table_chunks.append(" | ".join(cells))
        return "\n\n".join(paragraphs + table_chunks)

    def _extract_pdf(self, data: bytes) -> str:
        reader = PdfReader(io.BytesIO(data))
        pages = []
        for page in reader.pages:
            page_text = page.extract_text() or ""
            if page_text.strip():
                pages.append(page_text.strip())
        return "\n\n".join(pages)

    def _ocr_pdf(self, data: bytes) -> str:
        try:
            import pytesseract
            from pdf2image import convert_from_bytes
        except ImportError as exc:  # pragma: no cover - optional dependency path
            LOGGER.warning("OCR dependencies missing: %s", exc)
            return ""
        images = convert_from_bytes(data)
        texts = [pytesseract.image_to_string(image, lang=self.config.ocr_language) for image in images]
        return "\n\n".join(texts)
